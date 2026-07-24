import os
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
from config import OUTPUT_DIR, GRADE_ORDER, PATH_NAMES


CN_FONT = "宋体"
EN_FONT = "Times New Roman"
BODY_SIZE = Pt(10.5)
TABLE_HEADER_SIZE = Pt(10)
TABLE_BODY_SIZE = Pt(9)


def _set_run_font(run, cn_font=CN_FONT, en_font=EN_FONT, size=None, bold=False):
    run.font.name = en_font
    run.bold = bold
    if size:
        run.font.size = size
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)


def _set_paragraph_font(paragraph, cn_font=CN_FONT, en_font=EN_FONT, size=BODY_SIZE):
    for run in paragraph.runs:
        _set_run_font(run, cn_font, en_font, size)


def _setup_chinese_font():
    font_paths = [
        "C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc", "C:/Windows/Fonts/msyhbd.ttc",
    ]
    for fp in font_paths:
        if os.path.exists(fp):
            fm.fontManager.addfont(fp)
            prop = fm.FontProperties(fname=fp)
            plt.rcParams["font.family"] = prop.get_name()
            plt.rcParams["axes.unicode_minus"] = False
            return
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei"]
    plt.rcParams["axes.unicode_minus"] = False


def _save_fig(fig, filename):
    filepath = os.path.join(OUTPUT_DIR, filename)
    fig.savefig(filepath, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return filepath


def _gen_grade_pie(df, fname="grade_pie.png"):
    if "等级" not in df.columns:
        return None
    _setup_chinese_font()
    counts = df["等级"].value_counts().reindex(GRADE_ORDER).dropna()
    colors = ["#2ecc71", "#3498db", "#f39c12", "#e67e22", "#e74c3c", "#95a5a6"][:len(counts)]
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.pie(counts.values, labels=counts.index, colors=colors, autopct="%1.1f%%", startangle=90, pctdistance=0.75)
    ax.set_title("等级分布", fontsize=14, fontweight="bold")
    return _save_fig(fig, fname)


def _gen_path_bar(df, fname="path_bar.png"):
    if "路径" not in df.columns and "出海路径" not in df.columns:
        return None
    path_col = "出海路径" if "出海路径" in df.columns else "路径"
    _setup_chinese_font()
    counts = df[path_col].value_counts().sort_values(ascending=True)
    fig, ax = plt.subplots(figsize=(8, max(3, len(counts) * 0.4)))
    bars = ax.barh(counts.index, counts.values, color="#3498db", edgecolor="white")
    for bar, val in zip(bars, counts.values):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2, str(val), va="center", fontsize=9)
    ax.set_xlabel("企业数量")
    ax.set_title("路径分布", fontsize=14, fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    return _save_fig(fig, fname)


def _gen_score_box(df, fname="score_box.png"):
    score_cols = ["外贸基础能力", "电商运营能力", "合作配合意愿", "产能承接配套"]
    available = [c for c in score_cols if c in df.columns]
    if not available:
        return None
    _setup_chinese_font()
    valid_labels = []
    valid_data = []
    for c in available:
        vals = pd.to_numeric(df[c], errors="coerce").dropna().astype(float).tolist()
        if len(vals) >= 2:
            valid_labels.append(c)
            valid_data.append(vals)
    if not valid_data:
        return None
    fig, ax = plt.subplots(figsize=(8, 4))
    bp = ax.boxplot(valid_data, labels=valid_labels, patch_artist=True)
    for patch, color in zip(bp["boxes"], ["#3498db", "#2ecc71", "#e74c3c", "#f39c12"][:len(valid_labels)]):
        patch.set_facecolor(color)
        patch.set_alpha(0.6)
    ax.set_ylabel("得分")
    ax.set_title("各评分维度分布", fontsize=14, fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    return _save_fig(fig, fname)


def _gen_path_grade_stack(df, fname="path_grade_stack.png"):
    path_col = "出海路径" if "出海路径" in df.columns else "路径"
    if path_col not in df.columns or "等级" not in df.columns:
        return None
    ct = pd.crosstab(df[path_col], df["等级"])
    ct = ct.reindex(columns=[g for g in GRADE_ORDER if g in ct.columns], fill_value=0)
    if ct.empty:
        return None
    _setup_chinese_font()
    fig, ax = plt.subplots(figsize=(10, max(4, len(ct) * 0.5)))
    colors = ["#2ecc71", "#3498db", "#f39c12", "#e67e22", "#e74c3c", "#95a5a6"][:len(ct.columns)]
    ct.plot(kind="barh", stacked=True, ax=ax, color=colors, edgecolor="white", linewidth=0.5)
    ax.set_xlabel("企业数量")
    ax.set_title("路径×等级分布", fontsize=14, fontweight="bold")
    ax.legend(title="等级", bbox_to_anchor=(1.02, 1), loc="upper left", fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    return _save_fig(fig, fname)


def _gen_score_by_path(df, fname="score_by_path.png"):
    path_col = "出海路径" if "出海路径" in df.columns else "路径"
    score_cols = ["外贸基础能力", "电商运营能力", "合作配合意愿", "产能承接配套"]
    available = [c for c in score_cols if c in df.columns]
    if path_col not in df.columns or not available:
        return None
    grouped = df.groupby(path_col)[available].mean(numeric_only=True).dropna(how="all")
    if grouped.empty:
        return None
    grouped = grouped.dropna(axis=1, how="all")
    if grouped.empty or grouped.shape[1] == 0:
        return None
    available = list(grouped.columns)
    _setup_chinese_font()
    fig, ax = plt.subplots(figsize=(10, max(4, len(grouped) * 0.5)))
    grouped.plot(kind="barh", ax=ax, color=["#3498db", "#2ecc71", "#e74c3c", "#f39c12"][:len(available)], edgecolor="white", linewidth=0.5)
    ax.set_xlabel("平均得分")
    ax.set_title("各路径维度平均得分对比", fontsize=14, fontweight="bold")
    ax.legend(title="维度", bbox_to_anchor=(1.02, 1), loc="upper left", fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    return _save_fig(fig, fname)


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                _set_run_font(run, size=TABLE_HEADER_SIZE, bold=True)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    _set_run_font(run, size=TABLE_BODY_SIZE)
    return table


def generate_report(df, title="跨境电商企业分析报告", region_name="", mode="已有多阶段数据"):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    title_text = f"{region_name} {title}" if region_name else title
    heading = doc.add_heading(title_text, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        _set_run_font(run, size=Pt(22), bold=True)

    date_p = doc.add_paragraph(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(date_p)
    doc.add_paragraph()

    total = len(df)

    h = doc.add_heading("一、总览摘要", level=1)
    for run in h.runs:
        _set_run_font(run, size=Pt(14), bold=True)
    p = doc.add_paragraph(f"本次分析共涵盖 {total} 家企业。")
    _set_paragraph_font(p)

    if "等级" in df.columns:
        a_count = len(df[df["等级"] == "A级"])
        b_count = len(df[df["等级"] == "B级"])
        p = doc.add_paragraph(f"A级企业 {a_count} 家（{a_count/total*100:.1f}%），B级企业 {b_count} 家（{b_count/total*100:.1f}%）。")
        _set_paragraph_font(p)

    if "总分" in df.columns and df["总分"].notna().any():
        p = doc.add_paragraph(f"企业平均总分为 {df['总分'].mean():.1f} 分。")
        _set_paragraph_font(p)

    grade_img = None
    try:
        grade_img = _gen_grade_pie(df, f"{region_name}_grade.png")
    except Exception:
        pass
    if grade_img:
        doc.add_picture(grade_img, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h = doc.add_heading("二、等级分布详情", level=1)
    for run in h.runs:
        _set_run_font(run, size=Pt(14), bold=True)
    if "等级" in df.columns:
        headers = ["等级", "企业数量", "占比"]
        rows = []
        for grade in GRADE_ORDER:
            count = len(df[df["等级"] == grade])
            rows.append([grade, str(count), f"{count/total*100:.1f}%"])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    h = doc.add_heading("三、路径分布详情", level=1)
    for run in h.runs:
        _set_run_font(run, size=Pt(14), bold=True)
    path_col = "出海路径" if "出海路径" in df.columns else "路径"
    if path_col in df.columns:
        path_counts = df[path_col].value_counts()
        headers = ["路径", "路径名称", "企业数量", "占比"]
        rows = []
        for path, count in path_counts.items():
            if pd.notna(path) and str(path).strip() not in ["", "nan"]:
                rows.append([str(path), PATH_NAMES.get(str(path), ""), str(count), f"{count/total*100:.1f}%"])
        if rows:
            _add_table(doc, headers, rows)
            doc.add_paragraph()

        path_img = None
        try:
            path_img = _gen_path_bar(df, f"{region_name}_path.png")
        except Exception:
            pass
        if path_img:
            doc.add_picture(path_img, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h = doc.add_heading("四、路径×等级交叉分析", level=1)
    for run in h.runs:
        _set_run_font(run, size=Pt(14), bold=True)
    path_col = "出海路径" if "出海路径" in df.columns else "路径"
    if path_col in df.columns and "等级" in df.columns:
        ct = pd.crosstab(df[path_col], df["等级"])
        ct = ct.reindex(columns=[g for g in GRADE_ORDER if g in ct.columns], fill_value=0)
        if not ct.empty:
            headers = ["路径"] + [g for g in GRADE_ORDER if g in ct.columns]
            rows = []
            for path_name, row in ct.iterrows():
                rows.append([str(path_name)] + [str(int(v)) for v in row.values])
            _add_table(doc, headers, rows)
            doc.add_paragraph()

            stack_img = None
            try:
                stack_img = _gen_path_grade_stack(df, f"{region_name}_path_grade.png")
            except Exception:
                pass
            if stack_img:
                doc.add_picture(stack_img, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h = doc.add_heading("五、各路径企业得分对比", level=1)
    for run in h.runs:
        _set_run_font(run, size=Pt(14), bold=True)
    score_cols = ["外贸基础能力", "电商运营能力", "合作配合意愿", "产能承接配套"]
    available_dims = [c for c in score_cols if c in df.columns]
    if path_col in df.columns and available_dims:
        avg_df = df.groupby(path_col)[available_dims].mean(numeric_only=True).dropna(how="all")
        if not avg_df.empty:
            headers = ["路径"] + list(avg_df.columns)
            rows = []
            for path_name, row in avg_df.iterrows():
                rows.append([str(path_name)] + [f"{v:.1f}" if pd.notna(v) else "" for v in row.values])
            _add_table(doc, headers, rows)
            doc.add_paragraph()

            score_by_path_img = None
            try:
                score_by_path_img = _gen_score_by_path(df, f"{region_name}_score_by_path.png")
            except Exception:
                pass
            if score_by_path_img:
                doc.add_picture(score_by_path_img, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h = doc.add_heading("六、各维度评分分析", level=1)
    for run in h.runs:
        _set_run_font(run, size=Pt(14), bold=True)
    score_cols = {"外贸基础能力": 25, "电商运营能力": 23, "合作配合意愿": 35, "产能承接配套": 17}
    available_dims = {k: v for k, v in score_cols.items() if k in df.columns}
    if available_dims:
        headers = ["维度", "满分", "平均分", "最高分", "最低分", "中位数"]
        rows = []
        for dim, max_val in available_dims.items():
            col_data = df[dim].dropna()
            if len(col_data) > 0:
                rows.append([dim, str(max_val), f"{col_data.mean():.1f}", f"{col_data.max():.0f}", f"{col_data.min():.0f}", f"{col_data.median():.1f}"])
        if rows:
            _add_table(doc, headers, rows)
            doc.add_paragraph()

        box_img = None
        try:
            box_img = _gen_score_box(df, f"{region_name}_score.png")
        except Exception:
            pass
        if box_img:
            doc.add_picture(box_img, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h = doc.add_heading("七、AB级企业明细", level=1)
    for run in h.runs:
        _set_run_font(run, size=Pt(14), bold=True)
    if "等级" in df.columns:
        ab_df = df[df["等级"].isin(["A级", "B级"])].copy()
        if "总分" in ab_df.columns:
            ab_df = ab_df.sort_values("总分", ascending=False)

        display_cols = ["企业名称", "等级"]
        if path_col in ab_df.columns:
            display_cols.append(path_col)
        if "总分" in ab_df.columns:
            display_cols.append("总分")
        for extra in ["所属地区", "产业集群", "主营产品类别", "联系人", "联系电话"]:
            if extra in ab_df.columns:
                display_cols.append(extra)

        available_display = [c for c in display_cols if c in ab_df.columns]
        if len(ab_df) > 0 and available_display:
            headers = available_display
            rows = []
            for _, row in ab_df.iterrows():
                rows.append([str(row[c]) if pd.notna(row[c]) else "" for c in available_display])
            _add_table(doc, headers, rows)
        else:
            p = doc.add_paragraph("无AB级企业数据")
            _set_paragraph_font(p)

    output_path = os.path.join(OUTPUT_DIR, f"{region_name}_分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(output_path)
    return output_path
